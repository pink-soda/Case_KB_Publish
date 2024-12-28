'''
Author: pink-soda luckyli0127@gmail.com
Date: 2024-11-25 23:26:46
LastEditors: pink-soda luckyli0127@gmail.com
LastEditTime: 2024-12-19 14:21:30
FilePath: \undefinede:\Case_KB\trans.py
Description: 这是默认设置,请设置`customMade`, 打开koroFileHeader查看配置 进行设置: https://github.com/OBKoro1/koro1FileHeader/wiki/%E9%85%8D%E7%BD%AE
'''
import os
import uuid
from docx import Document
from PIL import Image
import pytesseract
import requests
from io import BytesIO

# 配置 Tesseract 的路径
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def ocr_image(image_data):
    """使用 OCR 识别图片中的文字"""
    with Image.open(BytesIO(image_data)) as img:
        return pytesseract.image_to_string(img)

def replace_images_with_text(doc_path, output_path):
    """读取 Word 文档，替换图片为文字，并保存新文档"""
    doc = Document(doc_path)
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            # 检查目标是否是外部链接（通过 URL 结构判断）
            if rel.target_ref.startswith("http://") or rel.target_ref.startswith("https://"):
                # 下载外部链接图片
                try:
                    response = requests.get(rel.target_ref)
                    response.raise_for_status()  # 检查 HTTP 状态码
                    image_data = response.content
                except Exception as e:
                    print(f"Failed to download image from {rel.target_ref}: {e}")
                    text = "[External image could not be downloaded]"
                    continue
            else:
                # 嵌入式图片
                image_data = rel.target_part.blob
            
            # 执行 OCR 识别
            try:
                text = ocr_image(image_data)
            except Exception as e:
                print(f"Error processing image: {e}")
                text = "[Image could not be processed]"
            
            # 替换文档中对应的段落
            for paragraph in doc.paragraphs:
                if rel_id in paragraph._p.xml:
                    paragraph.clear()  # 清空段落内容
                    paragraph.add_run(text)  # 添加识别的文字
                    break

    # 保存修改后的文档
    doc.save(output_path)

# 使用示例
input_doc = 'mail.docx'  # 输入文档的路径
output_doc = 'output.docx'  # 输出文档的路径
replace_images_with_text(input_doc, output_doc)
